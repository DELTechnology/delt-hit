#!/usr/bin/env python3
"""Convert Martinelli_etal_manuscript.docx to LaTeX."""

import os
import re
import docx
from docx.oxml.ns import qn

PAPER_DIR = "/Users/adrianomartinelli/projects/delt-hit/paper"
TABLES_DIR = os.path.join(PAPER_DIR, "tables")
os.makedirs(TABLES_DIR, exist_ok=True)

doc = docx.Document(os.path.join(PAPER_DIR, "Martinelli_etal_manuscript.docx"))


def escape_latex(text):
    """Escape special LaTeX characters."""
    # Order matters - backslash first
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"),
        ("%", r"\%"),
        ("#", r"\#"),
        ("$", r"\$"),
        ("^", r"\^{}"),
        ("~", r"\textasciitilde{}"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("_", r"\_"),
        ("<", r"\textless{}"),
        (">", r"\textgreater{}"),
        ("|", r"\textbar{}"),
    ]
    # Fix double-escape of textbackslash
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def escape_cell(text):
    """Escape cell text for LaTeX table."""
    text = text.strip()
    return escape_latex(text)


def table_to_latex(table, table_num, caption=""):
    """Convert a docx table to a LaTeX table environment."""
    rows = table.rows
    if not rows:
        return ""

    num_cols = max(len(row.cells) for row in rows)
    col_spec = "l" * num_cols

    lines = []
    lines.append(r"\begin{table}[htbp]")
    lines.append(r"\centering")
    if caption:
        lines.append(f"\\caption{{{escape_latex(caption)}}}")
    lines.append(f"\\label{{tab:table{table_num}}}")
    lines.append(f"\\begin{{tabular}}{{{col_spec}}}")
    lines.append(r"\toprule")

    for ri, row in enumerate(rows):
        cells = []
        seen_text = set()
        prev_text = None
        cell_texts = []
        # Handle merged cells by tracking unique cells
        actual_cells = row.cells
        # Deduplicate merged cells
        unique_cells = []
        seen_ids = []
        for cell in actual_cells:
            cid = cell._tc
            if cid not in seen_ids:
                seen_ids.append(cid)
                unique_cells.append(cell)

        for cell in unique_cells:
            txt = cell.text.replace("\n", " ").strip()
            cells.append(escape_cell(txt))

        # Pad to num_cols if needed (shouldn't usually happen)
        while len(cells) < num_cols:
            cells.append("")

        line = " & ".join(cells) + r" \\"
        lines.append(line)

        if ri == 0:
            lines.append(r"\midrule")

    lines.append(r"\bottomrule")
    lines.append(r"\end{tabular}")
    lines.append(r"\end{table}")

    return "\n".join(lines)


# Table captions from document analysis
TABLE_CAPTIONS = {
    1: "Input file specifications and sources",
    2: "Tutorial datasets for protocol validation",
    3: "Experiment configuration parameters",
    4: "Selection experimental design",
    5: "DNA sequence structure definition",
    6: "Constant DNA region sequences",
    7: "Building block definition sheet (example for B0)",
    8: "Reaction SMIRKS templates",
    9: "Compound definitions",
    10: "Additional reaction step definitions",
    11: "Additional reaction step definitions (special case)",
    12: "Hit identification by edgeR. Shown compounds are for FDR<0.05 and sorted by logFC to identify the top hits.",
    13: "Common issues and solutions",
    14: "Timing estimates for different dataset sizes",
}

# Write table files
for i, table in enumerate(doc.tables):
    table_num = i + 1
    caption = TABLE_CAPTIONS.get(table_num, f"Table {table_num}")
    latex = table_to_latex(table, table_num, caption)
    out_path = os.path.join(TABLES_DIR, f"table{table_num}.tex")
    with open(out_path, "w") as f:
        f.write(latex)
    print(f"Written {out_path}")


# =============================================
# Now build main.tex
# =============================================

paragraphs = doc.paragraphs

# Figure captions (paragraph index -> figure number)
FIGURE_PARA_TO_NUM = {
    39: 1,
    266: 2,
    285: 3,
    339: 4,
    340: 5,
    365: 6,
    420: 7,
}
# Figure 8 - need to check around para 417
# Let's scan for Figure 8
for i, p in enumerate(paragraphs):
    if "Figure 8" in p.text or "figure 8" in p.text.lower():
        print(f"Found Figure 8 mention at para [{i}]: {repr(p.text[:80])}")

# Table positions (paragraph index -> table number)
TABLE_PARA_TO_NUM = {
    130: 1,
    139: 2,
    174: 3,
    178: 4,
    196: 5,
    205: 6,
    211: 7,
    224: 8,
    228: 9,
    232: 10,
    235: 11,
    422: 12,
    425: 13,
    448: 14,
}

# Track which tables have been inserted
tables_inserted = set()
figures_inserted = set()

# Figure 8 - scan the document
fig8_idx = None
for i, p in enumerate(paragraphs):
    if re.search(r'[Ff]igure\s+8', p.text):
        print(f"Fig8 candidate [{i}]: {repr(p.text[:100])}")
        if fig8_idx is None:
            fig8_idx = i

print(f"Figure 8 paragraph: {fig8_idx}")

# Collect full figure captions from document
FIGURE_CAPTIONS = {}
for pi, pnum in FIGURE_PARA_TO_NUM.items():
    p = paragraphs[pi]
    # Strip "Figure N | " prefix
    text = p.text.strip()
    m = re.match(r'Figure\s+\d+\s*\|?\s*(.*)', text, re.DOTALL)
    if m:
        FIGURE_CAPTIONS[pnum] = m.group(1).strip()
    else:
        FIGURE_CAPTIONS[pnum] = text

if fig8_idx:
    text = paragraphs[fig8_idx].text.strip()
    m = re.match(r'Figure\s+8\s*\|?\s*(.*)', text, re.DOTALL)
    FIGURE_CAPTIONS[8] = m.group(1).strip() if m else text
    FIGURE_PARA_TO_NUM[fig8_idx] = 8


def make_figure(fig_num, caption):
    cap_escaped = escape_latex(caption)
    return f"""
\\begin{{figure}}[htbp]
\\centering
\\includegraphics[width=\\textwidth]{{figures/image{fig_num}.png}}
\\caption{{{cap_escaped}}}
\\label{{fig:figure{fig_num}}}
\\end{{figure}}
"""


def para_to_latex(p, idx):
    """Convert a paragraph to LaTeX."""
    style = p.style.name
    text = p.text.strip()

    # Skip empty paragraphs (but not all - some are spacers)
    if not text:
        return ""

    # Check for figure caption paragraphs
    if idx in FIGURE_PARA_TO_NUM:
        fig_num = FIGURE_PARA_TO_NUM[idx]
        if fig_num not in figures_inserted:
            figures_inserted.add(fig_num)
            caption = FIGURE_CAPTIONS.get(fig_num, text)
            return make_figure(fig_num, caption)
        return ""

    # Check for table caption paragraphs
    if idx in TABLE_PARA_TO_NUM:
        tnum = TABLE_PARA_TO_NUM[idx]
        if tnum not in tables_inserted:
            tables_inserted.add(tnum)
            return f"\n\\input{{tables/table{tnum}}}\n"
        return ""

    # Headings
    if style == "Heading 1":
        return f"\n\\section*{{{escape_latex(text)}}}\n"
    elif style == "Heading 2":
        return f"\n\\subsection{{{escape_latex(text)}}}\n"
    elif style == "Heading 3":
        return f"\n\\subsubsection{{{escape_latex(text)}}}\n"
    elif style == "Heading 4":
        return f"\n\\paragraph{{{escape_latex(text)}}}\n"
    elif style == "Heading 5":
        return f"\n\\subparagraph{{{escape_latex(text)}}}\n"

    # Code blocks
    if style == "code":
        # Don't escape code content
        return f"\n\\begin{{lstlisting}}\n{p.text}\n\\end{{lstlisting}}\n"

    # List items
    if style in ("List Paragraph", "Compact", "sheet"):
        return f"\\item {escape_latex(text)}"

    # Figure captions (non-table, non-figure ones)
    if style == "figure-caption":
        # Already handled above for actual figures/tables
        # This handles other figure-caption styled paragraphs
        if text:
            return f"\\textit{{{escape_latex(text)}}}\n\n"
        return ""

    # Bibliography
    if style == "EndNote Bibliography":
        # Strip leading number+tab
        clean = re.sub(r'^\d+\.\s*', '', text)
        return f"\\bibitem{{ref}} {escape_latex(clean)}\n"

    # Normal/body text
    return f"{escape_latex(text)}\n\n"


# Now build the document, handling list grouping
def build_body():
    lines = []
    i = 0
    n = len(paragraphs)

    # Skip title (para 0) and authors block (paras 1-6) - handle separately
    # Start from para 7

    in_list = False
    in_biblio = False

    while i < n:
        p = paragraphs[i]
        style = p.style.name
        text = p.text.strip()

        # Skip Authors heading/block and Abstract heading/text (handled in preamble)
        if i in (0, 1, 2, 3, 4, 5, 6, 7, 8, 9):
            i += 1
            continue

        # Handle bibliography section
        if style == "EndNote Bibliography":
            if not in_biblio:
                if in_list:
                    lines.append("\\end{itemize}\n")
                    in_list = False
                lines.append("\\begin{thebibliography}{99}\n")
                in_biblio = True
            clean = re.sub(r'^\d+\.\s*', '', p.text.strip())
            lines.append(f"\\bibitem{{ref{i}}} {escape_latex(clean)}\n")
            i += 1
            continue

        if in_biblio and style != "EndNote Bibliography":
            lines.append("\\end{thebibliography}\n")
            in_biblio = False

        # List items
        if style in ("List Paragraph", "Compact", "sheet"):
            if not in_list:
                lines.append("\\begin{itemize}\n")
                in_list = True
            if text:
                lines.append(f"  \\item {escape_latex(text)}\n")
            i += 1
            continue
        else:
            if in_list:
                lines.append("\\end{itemize}\n")
                in_list = False

        # Figure captions
        if i in FIGURE_PARA_TO_NUM:
            fig_num = FIGURE_PARA_TO_NUM[i]
            if fig_num not in figures_inserted:
                figures_inserted.add(fig_num)
                caption = FIGURE_CAPTIONS.get(fig_num, text)
                lines.append(make_figure(fig_num, caption))
            i += 1
            continue

        # Table captions
        if i in TABLE_PARA_TO_NUM:
            tnum = TABLE_PARA_TO_NUM[i]
            if tnum not in tables_inserted:
                tables_inserted.add(tnum)
                lines.append(f"\n\\input{{tables/table{tnum}}}\n")
            i += 1
            continue

        # Headings
        if style == "Heading 1":
            lines.append(f"\n\\section*{{{escape_latex(text)}}}\n")
            i += 1
            continue
        elif style == "Heading 2":
            if text:
                lines.append(f"\n\\subsection{{{escape_latex(text)}}}\n")
            i += 1
            continue
        elif style == "Heading 3":
            if text:
                lines.append(f"\n\\subsubsection{{{escape_latex(text)}}}\n")
            i += 1
            continue
        elif style == "Heading 4":
            if text:
                lines.append(f"\n\\paragraph{{{escape_latex(text)}}}\n\n")
            i += 1
            continue
        elif style == "Heading 5":
            if text:
                lines.append(f"\n\\subparagraph{{{escape_latex(text)}}}\n\n")
            i += 1
            continue

        # Code blocks
        if style == "code":
            # Collect consecutive code paragraphs
            code_lines = [p.text]
            j = i + 1
            while j < n and paragraphs[j].style.name == "code":
                code_lines.append(paragraphs[j].text)
                j += 1
            code_content = "\n".join(code_lines)
            # Replace Unicode chars with ASCII equivalents for lstlisting
            box_replacements = [
                ("├", "|--"),  # ├
                ("└", "`--"),  # └
                ("─", "-"),    # ─
                ("│", "|"),    # │
                ("•", "*"),    # •
                ("'", "'"),    # '
                ("'", "'"),    # '
                (""", '"'),    # "
                (""", '"'),    # "
                ("–", "--"),   # –
                ("—", "---"),  # —
                ("²", "^2"),   # ²
                ("³", "^3"),   # ³
                ("′", "'"),    # ′
                ("°", "deg"),  # °
                ("…", "..."),  # …
                (" ", " "),    # non-breaking space
            ]
            for uc, asc in box_replacements:
                code_content = code_content.replace(uc, asc)
            # Replace any remaining non-ASCII with ?
            code_content = code_content.encode('ascii', errors='replace').decode('ascii')
            lines.append(f"\n\\begin{{lstlisting}}\n{code_content}\n\\end{{lstlisting}}\n")
            i = j
            continue


        # figure-caption style
        if style == "figure-caption":
            if text and not text.startswith("Table") and not text.startswith("Figure"):
                lines.append(f"\\textit{{{escape_latex(text)}}}\n\n")
            i += 1
            continue

        # Normal text, Body Text, First Paragraph, etc.
        if text:
            lines.append(f"{escape_latex(text)}\n\n")

        i += 1

    if in_list:
        lines.append("\\end{itemize}\n")
    if in_biblio:
        lines.append("\\end{thebibliography}\n")

    return "".join(lines)


body = build_body()

# Build the abstract text
abstract_text = paragraphs[9].text.strip()

# Author info
author_line = re.sub(r'\s*\n\s*', ", ", paragraphs[2].text).strip().rstrip(",")
affil1 = paragraphs[3].text.replace("\n", " ").strip()
affil_footnotes = "; ".join(p.text.replace("\n", " ").strip() for p in [paragraphs[4], paragraphs[5], paragraphs[6]] if p.text.strip())

title = paragraphs[0].text.strip()

PREAMBLE = r"""\documentclass[11pt,a4paper]{article}

\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage[margin=2.5cm]{geometry}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage{hyperref}
\usepackage{amsmath}
\usepackage{longtable}
\usepackage{array}
\usepackage{caption}
\usepackage{subcaption}
\usepackage{xcolor}
\usepackage{listings}
\usepackage{microtype}
\usepackage{parskip}

\lstset{
  basicstyle=\small\ttfamily,
  breaklines=true,
  breakatwhitespace=false,
  frame=single,
  backgroundcolor=\color{gray!10},
  xleftmargin=1em,
  xrightmargin=1em,
}

\hypersetup{
  colorlinks=true,
  linkcolor=blue,
  urlcolor=blue,
  citecolor=blue,
}

"""

title_block = f"""\\title{{{escape_latex(title)}}}
\\author{{{escape_latex(author_line)}\\\\
\\small {escape_latex(affil1)}\\\\
\\small {escape_latex(affil_footnotes)}}}
\\date{{}}

\\begin{{document}}
\\maketitle

\\begin{{abstract}}
{escape_latex(abstract_text)}
\\end{{abstract}}

"""

main_tex = PREAMBLE + title_block + body + "\n\\end{document}\n"

out_main = os.path.join(PAPER_DIR, "main.tex")
with open(out_main, "w") as f:
    f.write(main_tex)

print(f"\nWritten {out_main}")
print("Done!")
